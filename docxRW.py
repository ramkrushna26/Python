
import docx

d = doc.Document("file.docx")
#get all paragraphs
print(d.paragraphs)

#get paragraph content
print(d.paragraphs[0].text)

#runs -  whenever change in style
print(p.runs)
print(p.runs[0].text)
#to check whether style is bold, will return True is style is set
print(p.runs[0].bold)
print(p.runs[0].underline)

#to assgin style
print(p.runs[0].bold = True)
print(p.runs[0].text)


#create new doc
nd = doc.Documet()
#add paragraph
nd.add_paragraph("add para content here")
nd.add_paragraph("add para content here")
#add runs
p = nd.paragraph[0]
print(p.runs)
p.runs[0].bold =True

nd.save('file.docx')

#get the all text of document
def getContent(doc1):
    doc = doc.Document(doc1)
    content = []
    for p in doc.paragraphs:
        content.append(p.text)
    return '\n'.join(content)

print(getContent('docx file'))
